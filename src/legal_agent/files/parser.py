from __future__ import annotations

from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import Any

from legal_agent.core.config import Settings
from legal_agent.core.ids import new_id


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".docx"}


@dataclass(frozen=True)
class ParsedUpload:
    file_id: str
    original_filename: str
    content_type: str | None
    size_bytes: int
    sha256_hex: str
    storage_path: str
    parse_status: str
    text_preview: str
    chunks: list[dict[str, Any]]
    metadata: dict[str, Any]


def parse_and_store_upload(
    settings: Settings,
    *,
    filename: str,
    content_type: str | None,
    data: bytes,
) -> ParsedUpload:
    original_filename = Path(filename or "upload.txt").name
    suffix = Path(original_filename).suffix.lower()
    file_id = new_id("file")
    digest = sha256(data).hexdigest()
    storage_path = _write_upload(settings, file_id, original_filename, data)
    parse_status = "PARSED"
    parse_error: str | None = None
    try:
        text = _extract_text(original_filename, data)
    except ValueError as exc:
        text = ""
        parse_status = "UNSUPPORTED"
        parse_error = str(exc)
    except Exception as exc:
        text = ""
        parse_status = "FAILED"
        parse_error = f"{type(exc).__name__}: {exc}"
    chunks = _chunks(file_id, original_filename, text) if text else []
    return ParsedUpload(
        file_id=file_id,
        original_filename=original_filename,
        content_type=content_type,
        size_bytes=len(data),
        sha256_hex=digest,
        storage_path=storage_path,
        parse_status=parse_status,
        text_preview=text[:500] if text else None,
        chunks=chunks,
        metadata={
            "extension": suffix,
            "parse_error": parse_error,
            "supported_extensions": sorted(SUPPORTED_EXTENSIONS),
        },
    )


def _write_upload(settings: Settings, file_id: str, filename: str, data: bytes) -> str:
    safe_name = "".join(ch if ch.isalnum() or ch in {".", "_", "-"} else "_" for ch in filename) or "upload.bin"
    path = settings.data_dir / "uploaded-files" / file_id / safe_name
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    return str(path)


def _extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return data.decode("utf-8", errors="replace")
    if suffix == ".docx":
        return _extract_docx_text(data)
    raise ValueError(f"unsupported upload extension: {suffix or '<none>'}")


def _extract_docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _chunks(file_id: str, filename: str, text: str, *, max_chars: int = 800) -> list[dict[str, Any]]:
    paragraphs = [part.strip() for part in text.replace("\r\n", "\n").split("\n\n") if part.strip()]
    if not paragraphs:
        paragraphs = [text.strip()] if text.strip() else []
    chunks: list[dict[str, Any]] = []
    current = ""
    for paragraph in paragraphs:
        if current and len(current) + len(paragraph) + 2 > max_chars:
            chunks.append(_chunk(file_id, filename, len(chunks), current))
            current = paragraph
        else:
            current = paragraph if not current else f"{current}\n\n{paragraph}"
    if current:
        chunks.append(_chunk(file_id, filename, len(chunks), current))
    return chunks


def _chunk(file_id: str, filename: str, index: int, content: str) -> dict[str, Any]:
    return {
        "chunk_id": f"{file_id}_chunk_{index + 1:03d}",
        "file_id": file_id,
        "chunk_index": index,
        "page_no": None,
        "content": content,
        "citation_anchor": f"{filename}#chunk-{index + 1}",
        "metadata": {"filename": filename},
    }

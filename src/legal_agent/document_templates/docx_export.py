from __future__ import annotations

import re
from pathlib import Path


NUMBERED_RE = re.compile(r"^\d+\.\s+")


def write_markdown_docx(markdown: str, output_path: Path) -> None:
    from docx import Document

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "劳动人事争议仲裁申请书"

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line == "---":
            document.add_paragraph()
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if NUMBERED_RE.match(line):
            document.add_paragraph(NUMBERED_RE.sub("", line).strip(), style="List Number")
            continue
        paragraph = document.add_paragraph(line)
        if line.startswith("提示："):
            for run in paragraph.runs:
                run.italic = True

    document.save(output_path)
